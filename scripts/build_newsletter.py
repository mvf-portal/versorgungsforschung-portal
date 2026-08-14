#!/usr/bin/env python3
"""Erzeugt aus studien-archiv.json die Newsletter- und Download-Dateien.

Ausgabe:
  studien-feed.xml            RSS 2.0 fuer Mailchimps RSS-to-Email
  download/studien-aktuell.docx / .csv     nur der juengste Tag
  download/studien-archiv.docx  / .csv     der vollstaendige Bestand

Liest ausschliesslich das Archiv - kein API-Key, kein Netzzugriff. Laesst sich
deshalb jederzeit einzeln aufrufen, um alles aus dem Bestand neu aufzubauen:

    py scripts/build_newsletter.py

Die Feed-Logik im Kern: Ein <item> je Studie, die GUID ist die PMID. Mailchimp
verschickt nur Items, die es noch nicht gesendet hat - eine Studie, die an
mehreren Tagen in den PubMed-Treffern auftaucht, geht dadurch nie doppelt raus.
"""
from __future__ import annotations

import csv
import datetime as dt
import io
import json
import os
import sys
from email.utils import format_datetime
from xml.sax.saxutils import escape
from zoneinfo import ZoneInfo

ARCHIVE = "studien-archiv.json"
FEED = "studien-feed.xml"
DL_DIR = "download"

HUB = "https://wissen.m-vf.de/"
FEED_URL = "https://wissen.m-vf.de/studien-feed.xml"

# Wie viele Studien der Feed vorhaelt. Mailchimp braucht nur die noch nicht
# gesendeten; ein Puffer schuetzt vor Luecken, wenn ein Versand ausfaellt.
FEED_MAX = 60

TZ = ZoneInfo("Europe/Berlin")
# Die Aktualisierung laeuft um 06:00 Berliner Zeit; das Archiv haelt nur das
# Datum fest. Diese Uhrzeit macht daraus einen gueltigen RFC-822-Zeitstempel.
UPDATE_HOUR = 6


def utm(url: str, medium: str, content: str = "") -> str:
    """Haengt die Kampagnenparameter an - Newsletter-Klicks bleiben zurechenbar."""
    sep = "&" if "?" in url else "?"
    p = f"utm_source=newsletter&utm_medium={medium}&utm_campaign=studien-feed"
    if content:
        p += f"&utm_content={content}"
    return f"{url}{sep}{p}"


def load() -> list[dict]:
    with open(ARCHIVE, encoding="utf-8") as f:
        entries = json.load(f)
    if not entries:
        raise RuntimeError(f"{ARCHIVE} ist leer")
    # Neueste zuerst; innerhalb eines Tages stabil ueber die PMID.
    entries.sort(key=lambda e: (e["aufgenommen"], e["pmid"]), reverse=True)
    return entries


def pubdate(iso: str, rang: int) -> dt.datetime:
    """Datum + Uhrzeit fuer ein Feed-Item.

    Der Rang wird sekundenweise ABGEZOGEN, nicht addiert: Mailchimp sortiert die
    Items eines Versands nach pubDate: die erste Studie der Tagesauswahl braucht
    also den spaetesten Zeitstempel. Addieren wuerde die Tagesauswahl im
    Newsletter verkehrt herum ausgeben.
    """
    d = dt.date.fromisoformat(iso)
    basis = dt.datetime(d.year, d.month, d.day, UPDATE_HOUR, 0, 0, tzinfo=TZ)
    return basis - dt.timedelta(seconds=rang)


# --------------------------------------------------------------------- Feed

def item_html(e: dict) -> str:
    """Der Inhalt eines Feed-Items - bewusst schlicht.

    Outlook rendert mit der Word-Engine: keine CSS-Variablen, kein Flexbox,
    kein <style>-Block. Nur Inline-Styles auf simplen Block-Elementen.
    """
    pubmed = utm(f"https://pubmed.ncbi.nlm.nih.gov/{e['pmid']}/", "email", "studie")
    # Journal und Jahr stehen bereits im <title> des Items (und damit in jeder
    # Mailchimp-Vorlage ueber der Zusammenfassung) - hier nicht wiederholen.
    kopf = "  ·  ".join(x for x in (e.get("author"), e.get("pubdate")) if x)
    return (
        (f'<p style="margin:0 0 8px;font:italic 13px Georgia,serif;color:#666;">'
         f'{escape(kopf)}</p>' if kopf else '')
        + f'<p style="margin:0 0 10px;font:15px Georgia,serif;color:#333;">'
        f'{escape(e["sum"])}</p>'
        f'<table role="presentation" cellpadding="0" cellspacing="0" border="0" width="100%">'
        f'<tr><td style="background:#f4f6f4;border-left:3px solid #5a7d3a;padding:10px 14px;'
        f'font:15px Georgia,serif;color:#222;">'
        f'<strong>Ergebnis:</strong> {escape(e["result"])}</td></tr></table>'
        f'<p style="margin:10px 0 0;font:13px Georgia,serif;">'
        f'<a href="{escape(pubmed)}" style="color:#5a7d3a;">Studie in PubMed ansehen &rarr;</a></p>'
    )


def cdata(s: str) -> str:
    # ]]> waere ein vorzeitiges Ende des CDATA-Abschnitts.
    return "<![CDATA[" + s.replace("]]>", "]]&gt;") + "]]>"


def build_feed(entries: list[dict]) -> str:
    # Bewusst NICHT die aktuelle Uhrzeit: die Ausgabe soll allein vom Archiv
    # abhaengen. Sonst entstuende bei jedem Lauf ein Diff und damit auch an
    # Tagen ohne neue Studien ein Commit samt Pages-Build.
    stand = pubdate(entries[0]["aufgenommen"], 0)
    teile = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<rss version="2.0" xmlns:atom="http://www.w3.org/2005/Atom">',
        "<channel>",
        "<title>MVF-Knowledge-Hub: Neueste Studien der Versorgungsforschung</title>",
        f"<link>{escape(utm(HUB, 'email', 'kanal'))}</link>",
        "<description>Täglich ausgewählte Studien aus PubMed, auf Deutsch "
        "zusammengefasst mit den konkreten Ergebnissen. Ein Angebot von "
        "Monitor Versorgungsforschung.</description>",
        "<language>de-de</language>",
        f"<lastBuildDate>{format_datetime(stand)}</lastBuildDate>",
        f'<atom:link href="{escape(FEED_URL)}" rel="self" type="application/rss+xml"/>',
    ]
    for rang, e in enumerate(entries[:FEED_MAX]):
        titel = f'{e["title"]} ({e["journal"]} {e["year"]})'
        link = utm(f"https://pubmed.ncbi.nlm.nih.gov/{e['pmid']}/", "email", "studie")
        teile += [
            "<item>",
            f"<title>{escape(titel)}</title>",
            f"<link>{escape(link)}</link>",
            f'<guid isPermaLink="false">pmid-{escape(e["pmid"])}</guid>',
            f"<pubDate>{format_datetime(pubdate(e['aufgenommen'], rang))}</pubDate>",
            f"<description>{cdata(item_html(e))}</description>",
            "</item>",
        ]
    teile += ["</channel>", "</rss>", ""]
    return "\n".join(teile)


# ----------------------------------------------------------------- Downloads

KOPFTEXT = ("Ein Service der Knowledge-Datenbank von Monitor Versorgungsforschung. "
            "Täglich automatisiert KI-kuratiert aus PubMed.")
LOGO = "logo/mvf-logo.png"

SPALTEN = ["Aufgenommen", "Autor", "Publiziert am", "Journal", "Jahr", "Titel",
           "Fragestellung", "Ergebnis", "PMID", "PubMed-Link"]


def zeile(e: dict) -> list[str]:
    return [e["aufgenommen"], e.get("author", ""), e.get("pubdate", ""),
            e["journal"], e["year"], e["title"], e["sum"],
            e["result"], e["pmid"], f"https://pubmed.ncbi.nlm.nih.gov/{e['pmid']}/"]


def write_csv(pfad: str, entries: list[dict]) -> None:
    # Semikolon + BOM: so oeffnet Excel unter Windows die Datei korrekt in
    # Spalten und stellt Umlaute richtig dar.
    buf = io.StringIO(newline="")
    w = csv.writer(buf, delimiter=";", quoting=csv.QUOTE_MINIMAL, lineterminator="\r\n")
    # Eine Tabelle kann kein Logo aufnehmen - deshalb nur der Kopftext,
    # abgesetzt durch eine Leerzeile.
    w.writerow([KOPFTEXT])
    w.writerow([])
    w.writerow(SPALTEN)
    for e in entries:
        w.writerow(zeile(e))
    with open(pfad, "w", encoding="utf-8-sig", newline="") as f:
        f.write(buf.getvalue())


def normalize_docx(pfad: str, when: dt.datetime) -> None:
    """Schreibt das docx-ZIP mit festen Zeitstempeln neu.

    Ein docx ist ein ZIP-Archiv; python-docx traegt in jeden Eintrag die
    aktuelle Uhrzeit ein. Dadurch waere die Datei bei jedem Lauf binaer
    verschieden, es entstuende taeglich ein Commit samt Pages-Build - auch an
    Tagen ohne neue Studien. Reihenfolge und Kompression bleiben unveraendert,
    nur das Datum wird fixiert.
    """
    import zipfile

    fest = (when.year, when.month, when.day, when.hour, when.minute, when.second)
    with zipfile.ZipFile(pfad) as z:
        eintraege = [(i, z.read(i.filename)) for i in z.infolist()]
    with zipfile.ZipFile(pfad, "w", zipfile.ZIP_DEFLATED) as z:
        for info, daten in eintraege:
            neu = zipfile.ZipInfo(info.filename, date_time=fest)
            neu.compress_type = info.compress_type
            neu.external_attr = info.external_attr
            z.writestr(neu, daten)


def write_docx(pfad: str, entries: list[dict], titel: str, stand: str,
               zeitpunkt: dt.datetime) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    doc.core_properties.title = titel
    doc.core_properties.author = "Monitor Versorgungsforschung"
    # Auch die Dokument-Metadaten aus den Daten ableiten, nicht aus der Systemuhr.
    doc.core_properties.created = zeitpunkt.replace(tzinfo=None)
    doc.core_properties.modified = zeitpunkt.replace(tzinfo=None)
    doc.core_properties.revision = 1

    if os.path.exists(LOGO):
        doc.add_picture(LOGO, width=Inches(1.9))
    pk = doc.add_paragraph()
    rk = pk.add_run(KOPFTEXT)
    rk.font.size = Pt(9.5)
    rk.font.color.rgb = RGBColor(0x00, 0x51, 0xA1)

    doc.add_heading(titel, level=0)
    p = doc.add_paragraph()
    r = p.add_run(f"Stand: {stand}  |  {len(entries)} Studien  |  wissen.m-vf.de")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    tag = None
    for e in entries:
        if e["aufgenommen"] != tag:
            tag = e["aufgenommen"]
            d = dt.date.fromisoformat(tag)
            doc.add_heading(f"Aufgenommen am {d:%d.%m.%Y}", level=1)

        doc.add_heading(e["title"], level=2)

        q = doc.add_paragraph()
        teile = [x for x in (e.get("author"), e["journal"], e.get("pubdate") or e["year"]) if x]
        rq = q.add_run("  ·  ".join(teile) + f'  ·  PMID {e["pmid"]}')
        rq.italic = True
        rq.font.size = Pt(9)
        rq.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph(e["sum"])

        pe = doc.add_paragraph()
        pe.add_run("Ergebnis: ").bold = True
        pe.add_run(e["result"])

        pl = doc.add_paragraph()
        rl = pl.add_run(f'https://pubmed.ncbi.nlm.nih.gov/{e["pmid"]}/')
        rl.font.size = Pt(9)
        rl.font.color.rgb = RGBColor(0x5A, 0x7D, 0x3A)

    doc.save(pfad)
    normalize_docx(pfad, zeitpunkt)


def main() -> int:
    entries = load()
    # Ebenfalls aus dem Archiv abgeleitet, nicht aus der Systemuhr (s. build_feed).
    zeitpunkt = pubdate(entries[0]["aufgenommen"], 0)
    stand = f"{dt.date.fromisoformat(entries[0]['aufgenommen']):%d.%m.%Y}"

    with open(FEED, "w", encoding="utf-8", newline="\n") as f:
        f.write(build_feed(entries))
    print(f"{FEED}: {min(len(entries), FEED_MAX)} Items.")

    os.makedirs(DL_DIR, exist_ok=True)
    neuester = entries[0]["aufgenommen"]
    aktuell = [e for e in entries if e["aufgenommen"] == neuester]

    write_csv(f"{DL_DIR}/studien-aktuell.csv", aktuell)
    write_csv(f"{DL_DIR}/studien-archiv.csv", entries)
    write_docx(f"{DL_DIR}/studien-aktuell.docx", aktuell,
               "Neueste Studien der Versorgungsforschung", stand, zeitpunkt)
    write_docx(f"{DL_DIR}/studien-archiv.docx", entries,
               "Studienarchiv Versorgungsforschung", stand, zeitpunkt)
    print(f"{DL_DIR}/: aktuell {len(aktuell)} Studien, Archiv {len(entries)} Studien.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
